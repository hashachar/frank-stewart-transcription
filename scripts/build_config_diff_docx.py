#!/usr/bin/env python3
"""Build the multi-config comparison docx (cost + diacritic counts + word-level diff).

Extends the earlier 2-config `gpt-5.5_vs_5.6-sol_Phase1e_diff.docx` to N configs,
restricted to a fixed scan set so every config is compared on identical pages.

Sections:
  0  Cost analysis        — total / avg / 1200-scan projection, actual + repriced-at-flex
  A  Diacritic totals per scan        (interim-notation tokens, from the Phase 1 .txt)
  B  Breakdown by diacritic type      (ditto)
  C  Stacked diacritics                (base char carrying 2+ marks)
  D  Contextual word-level diff        (from the Phase 2 .literal.txt, Unicode)

Section D uses the FIRST config as the alignment base; each other config's text is
mapped onto the base's word indices via difflib opcodes, so all configs can be shown
side by side at the same difference site.
"""
import argparse
import datetime
import difflib
import re
import statistics
import unicodedata
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

import collect_config_costs as costs

BASE = Path(__file__).resolve().parent.parent

NAVY   = RGBColor(0x1A, 0x1A, 0x2E)
SLATE  = RGBColor(0x2C, 0x3E, 0x50)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREEN  = RGBColor(0x1E, 0x8A, 0x3C)
BLUE   = RGBColor(0x1F, 0x5C, 0xA8)
GREY   = RGBColor(0x88, 0x88, 0x88)

STRUCTURAL = {"LB", "HEADING", "/HEADING", "INDENT"}
TOKEN_RE = re.compile(r"\{([^{}]+)\}")
# a base char immediately followed by 2+ {mark} tokens = a stacked diacritic
STACK_RE = re.compile(r"([A-Za-z])((?:\{[^{}]+\}){2,})")

SCANS = [f"FN-{i:04d}" for i in range(1, 11)]


def phase1_txt(folder: Path, stem: str):
    c = [p for p in folder.glob(f"{stem}_*.txt")
         if ".literal" not in p.name and ".normalized" not in p.name]
    return sorted(c)[-1] if c else None


def literal_txt(folder: Path, stem: str):
    """Phase 2 literals live either beside the Phase 1 .txt (5.6-sol runs) or in a
    phase2/ subfolder (the older 5.5 run). Search both."""
    c = list(folder.glob(f"{stem}_*.literal.txt"))
    if not c:
        c = list((folder / "phase2").glob(f"{stem}_*.literal.txt"))
    return sorted(c)[-1] if c else None


def count_tokens(path: Path):
    """→ (Counter of diacritic tokens, total, Counter of stack combos)"""
    text = path.read_text(encoding="utf-8")
    counts = Counter()
    for m in TOKEN_RE.finditer(text):
        tok = m.group(1)
        if tok not in STRUCTURAL:
            counts[tok] += 1
    stacks = Counter()
    for m in STACK_RE.finditer(text):
        marks = TOKEN_RE.findall(m.group(2))
        marks = [x for x in marks if x not in STRUCTURAL]
        if len(marks) >= 2:
            stacks[" + ".join("{%s}" % x for x in marks)] += 1
    return counts, sum(counts.values()), stacks


def words_of(path: Path):
    txt = path.read_text(encoding="utf-8")
    txt = txt.replace("**", "")
    return txt.split()


def map_onto_base(base_words, other_words):
    """→ (mapping, insertions): base index → other words aligned to it."""
    sm = difflib.SequenceMatcher(a=base_words, b=other_words, autojunk=False)
    mapping = {k: [] for k in range(len(base_words))}
    insertions = {k: [] for k in range(len(base_words) + 1)}
    changed = set()
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for k in range(i1, i2):
                mapping[k] = [other_words[j1 + (k - i1)]]
        elif tag == "replace":
            for k in range(i1, i2):
                mapping[k] = []
                changed.add(k)
            if i1 < len(base_words):
                mapping[i1] = other_words[j1:j2]
            else:
                insertions[i1].extend(other_words[j1:j2])
        elif tag == "delete":
            for k in range(i1, i2):
                mapping[k] = []
                changed.add(k)
        elif tag == "insert":
            insertions[i1].extend(other_words[j1:j2])
            changed.add(min(i1, max(0, len(base_words) - 1)))
    return mapping, insertions, changed


def merge_regions(changed, n, pad=0):
    """Contiguous base-index runs from a set of changed indices."""
    if not changed:
        return []
    idx = sorted(changed)
    regions = []
    s = p = idx[0]
    for k in idx[1:]:
        if k <= p + 1 + pad:
            p = k
        else:
            regions.append((s, p + 1))
            s = p = k
    regions.append((s, p + 1))
    return regions


def render_region(cell, base_words, mapping, insertions, s, e, ctx=6, changed=None):
    """Write '…context DIFF context…' into a table cell."""
    p = cell.paragraphs[0]
    pre = " ".join(base_words[max(0, s - ctx):s])
    post = " ".join(base_words[e:e + ctx])
    mid = []
    for k in range(s, e):
        mid.extend(insertions.get(k, []))
        mid.extend(mapping.get(k, []))
    mid.extend(insertions.get(e, []))
    mid_text = " ".join(mid)

    if pre:
        r = p.add_run(("… " if s > ctx else "") + pre + " ")
        r.font.size = Pt(8); r.font.color.rgb = GREY
    r = p.add_run(mid_text if mid_text else "(omitted)")
    r.font.size = Pt(8); r.bold = True
    r.font.color.rgb = GREEN if mid_text else RED
    if not mid_text:
        r.italic = True
    if post:
        r = p.add_run(" " + post + (" …" if e + ctx < len(base_words) else ""))
        r.font.size = Pt(8); r.font.color.rgb = GREY


def richness(text: str) -> int:
    """How much diacritical detail a reading carries — combining marks plus
    non-ASCII letterforms (θ, ‘, ʿ, …). Used to mark the 'richest' variant at a
    difference site. Deliberately config-agnostic: it scores the TEXT, not who
    produced it. Richer ≠ correct — it can also mean over-detection."""
    if not text:
        return -1
    d = unicodedata.normalize("NFD", text)
    return sum(1 for ch in d if unicodedata.combining(ch)) + \
           sum(1 for ch in d if ord(ch) > 127 and not unicodedata.combining(ch))


def strip_marks(text: str) -> str:
    """Letterforms with all combining marks removed — 'Gṣāṛ.' → 'Gsar.'"""
    return "".join(ch for ch in unicodedata.normalize("NFD", text)
                   if not unicodedata.combining(ch))


def site_kind(variants) -> str:
    """Classify a disagreement: same letters, different marks = 'diacritic';
    different letters = 'wording'. Lets the reader separate 'did it see the
    macron?' from 'did it read the word?' — different questions."""
    return "diacritic" if len({strip_marks(v) for v in variants}) == 1 else "wording"


def variant_at(label, scaffold_label, base_words, maps, ins, s, e):
    """What `label` reads across base word-range [s,e)."""
    if label == scaffold_label:
        return " ".join(base_words[s:e])
    mapping, insertions = maps[label], ins[label]
    out = []
    for k in range(s, e):
        out.extend(insertions.get(k, []))
        out.extend(mapping.get(k, []))
    out.extend(insertions.get(e, []))
    return " ".join(out)


def add_table(doc, headers, rows, style="Light Grid Accent 1", size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(size)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(size)
            if i == 0:
                r.bold = True
    return t


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", required=True)
    ap.add_argument("--config", action="append", required=True,
                    metavar="LABEL=FOLDER",
                    help="Repeatable, in display order. First is the diff base.")
    args = ap.parse_args()

    configs = []
    for spec in args.config:
        label, folder = spec.split("=", 1)
        configs.append((label, BASE / "outputs" / folder))
    labels = [c[0] for c in configs]

    # ---------- gather ----------
    cost = {lab: costs.collect(folder, SCANS) for lab, folder in configs}
    counts, totals, stacks = {}, {}, {}
    for lab, folder in configs:
        counts[lab], totals[lab], stacks[lab] = {}, {}, Counter()
        for s in SCANS:
            p = phase1_txt(folder, s)
            if not p:
                continue
            c, tot, st = count_tokens(p)
            counts[lab][s] = c
            totals[lab][s] = tot
            stacks[lab] += st

    doc = Document()
    title = doc.add_heading("Bedouin Transcription — Config Comparison", level=0)
    title.runs[0].font.color.rgb = NAVY

    def meta(k, v):
        p = doc.add_paragraph()
        r = p.add_run(k); r.bold = True
        p.add_run(v)

    meta("Generated: ", f"{datetime.date.today():%Y-%m-%d}")
    meta("Scans: ", f"{SCANS[0]} – {SCANS[-1]} ({len(SCANS)}) — identical page set for every config")
    meta("Configs: ", "  ·  ".join(labels))
    meta("Common: ", "Code Interpreter on · Phase 1e prompt · /v1/responses")

    # ---------- Section 0: cost ----------
    h = doc.add_heading("0.  Cost analysis", level=1); h.runs[0].font.color.rgb = SLATE
    p = doc.add_paragraph()
    r = p.add_run("Every config is priced from its measured input/output tokens at Flex "
                  "rates (50% off standard), regardless of the tier it actually ran on. "
                  "This is the only apples-to-apples basis: the GPT-5.5 pages were run at "
                  "mixed tiers (default-tier sync and Batch) while both 5.6-sol runs were "
                  "Flex, so real historical spend is not comparable. Tokens are what the "
                  "model actually did; tier is just a price multiplier.")
    r.italic = True; r.font.size = Pt(9)

    rows = []
    proj = {}
    for lab in labels:
        d = [v for v in cost[lab].values() if "error" not in v]
        if not d:
            rows.append([lab, "—", "—", "—", "—", "—"])
            continue
        n = len(d)
        flx = sum(x["flex_cost"] for x in d)
        in_avg = sum(x["in"] for x in d) / n
        out_avg = sum(x["out"] for x in d) / n
        proj[lab] = flx / n * 1200
        rows.append([lab, f"{n}", f"{in_avg:,.0f}", f"{out_avg:,.0f}",
                     f"${flx:.2f}", f"${flx/n:.4f}"])
    add_table(doc, ["Config", "Scans", "Avg input tok", "Avg output tok",
                    "Total (10 scans)", "Avg / scan"], rows)

    doc.add_paragraph()
    add_table(doc, ["Config", "Projection to 1,200 scans (Flex rates)", "vs. baseline"],
              [[lab,
                f"${proj[lab]:,.0f}" if lab in proj else "—",
                ("baseline" if i == 0 else
                 (f"{(proj[lab]/proj[labels[0]]-1)*100:+.0f}%  (${proj[lab]-proj[labels[0]]:+,.0f})"
                  if lab in proj and labels[0] in proj else "—"))]
               for i, lab in enumerate(labels)])
    p = doc.add_paragraph()
    r = p.add_run("Projection = Flex average per scan × 1,200. It assumes the remaining "
                  "pages resemble FN-0001–0010 in density, and excludes retries, failed "
                  "requests, and Phase 2 (local, free). Flex 429 capacity rejections are "
                  "not billed and cost nothing.")
    r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY

    tier_note = "; ".join(
        f"{lab}: {'/'.join(sorted({x['tier'] for x in cost[lab].values() if 'error' not in x}))}"
        for lab in labels if any('error' not in x for x in cost[lab].values()))
    p = doc.add_paragraph()
    r = p.add_run(f"Tiers actually used (for reference only — not used in any figure above) — {tier_note}.")
    r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY

    # ---------- Section A ----------
    h = doc.add_heading("A.  Diacritic totals per scan", level=1); h.runs[0].font.color.rgb = SLATE
    p = doc.add_paragraph()
    r = p.add_run("Interim-notation tokens ({macron}, {dotbelow}, …); structural tokens "
                  "({LB}, {HEADING}, {INDENT}) excluded. More marks ≠ better — it can mean "
                  "finer detection or over-detection; read with Section D.")
    r.italic = True; r.font.size = Pt(9)
    rows = []
    for s in SCANS:
        row = [s] + [str(totals[l].get(s, "—")) for l in labels]
        b = totals[labels[0]].get(s)
        for l in labels[1:]:
            v = totals[l].get(s)
            row.append(f"{v-b:+d}" if (v is not None and b is not None) else "—")
        rows.append(row)
    tot_row = ["TOTAL"] + [str(sum(totals[l].values())) for l in labels]
    b_tot = sum(totals[labels[0]].values())
    for l in labels[1:]:
        tot_row.append(f"{sum(totals[l].values()) - b_tot:+d}")
    rows.append(tot_row)
    add_table(doc, ["Scan"] + labels + [f"Δ vs {labels[0]}" for l in labels[1:]], rows)

    # ---------- Section B ----------
    h = doc.add_heading("B.  Breakdown by diacritic type", level=1); h.runs[0].font.color.rgb = SLATE
    agg = {l: Counter() for l in labels}
    for l in labels:
        for s in SCANS:
            agg[l] += counts[l].get(s, Counter())
    types = sorted({k for l in labels for k in agg[l]},
                   key=lambda k: -agg[labels[0]].get(k, 0))
    rows = []
    for k in types:
        row = ["{%s}" % k] + [str(agg[l].get(k, 0)) for l in labels]
        b = agg[labels[0]].get(k, 0)
        for l in labels[1:]:
            row.append(f"{agg[l].get(k,0)-b:+d}")
        rows.append(row)
    add_table(doc, ["Diacritic type"] + labels + [f"Δ vs {labels[0]}" for l in labels[1:]], rows)

    # ---------- Section C ----------
    h = doc.add_heading("C.  Stacked diacritics (base char with 2+ marks)", level=1)
    h.runs[0].font.color.rgb = SLATE
    combos = sorted({k for l in labels for k in stacks[l]},
                    key=lambda k: -stacks[labels[0]].get(k, 0))
    rows = []
    for k in combos:
        row = [k] + [str(stacks[l].get(k, 0)) for l in labels]
        b = stacks[labels[0]].get(k, 0)
        for l in labels[1:]:
            row.append(f"{stacks[l].get(k,0)-b:+d}")
        rows.append(row)
    if rows:
        add_table(doc, ["Stack combination"] + labels + [f"Δ vs {labels[0]}" for l in labels[1:]], rows)
    p = doc.add_paragraph()
    for i, l in enumerate(labels):
        r = p.add_run(f"{l}: "); r.font.size = Pt(9)
        r = p.add_run(f"{sum(stacks[l].values())}"); r.bold = True; r.font.size = Pt(9)
        if i < len(labels) - 1:
            p.add_run("   ")

    # ---------- Section D ----------
    h = doc.add_heading("D.  Word-level disagreements (per scan)", level=1)
    h.runs[0].font.color.rgb = SLATE
    p = doc.add_paragraph()
    r = p.add_run("Each row is one site where the configs disagree. No config is treated "
                  "as the reference: every reading is shown next to the configs that "
                  "produced it, grouped so agreement is visible at a glance. ")
    r.italic = True; r.font.size = Pt(9)
    r = p.add_run("Where one reading carries strictly more diacritical detail than the "
                  "others it is marked ")
    r.italic = True; r.font.size = Pt(9)
    r = p.add_run("◆ richest"); r.italic = True; r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = GREEN
    r = p.add_run(". ")
    r.italic = True; r.font.size = Pt(9)
    r = p.add_run("Richest is NOT necessarily correct")
    r.italic = True; r.bold = True; r.font.size = Pt(9)
    r = p.add_run(" — it may equally be over-detection; the scan is the arbiter. Where "
                  "the readings tie on detail (e.g. a plain spelling difference) nothing "
                  "is marked, and the order carries no meaning.")
    r.italic = True; r.font.size = Pt(9)

    p = doc.add_paragraph()
    r = p.add_run("Each site is tagged ")
    r.italic = True; r.font.size = Pt(9)
    r = p.add_run("diacritic"); r.italic = True; r.bold = True; r.font.size = Pt(9)
    r = p.add_run(" (same letters, different marks — did the model see the macron?) or ")
    r.italic = True; r.font.size = Pt(9)
    r = p.add_run("wording"); r.italic = True; r.bold = True; r.font.size = Pt(9)
    r = p.add_run(" (different letters — did it read the word?). These are different "
                  "kinds of accuracy and are worth judging separately.")
    r.italic = True; r.font.size = Pt(9)

    p = doc.add_paragraph()
    r = p.add_run("Alignment ignores line breaks (the configs wrap lines differently, so "
                  "line numbers are not comparable) and matches on the word stream. The "
                  "internal alignment scaffold is the config of median length — an "
                  "indexing device only; it affects neither what is shown nor which "
                  "reading is favoured.")
    r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY

    grand = 0
    agree_all = 0
    dia_total = 0
    word_total = 0
    for s in SCANS:
        paths = {l: literal_txt(f, s) for l, f in configs}
        if any(v is None for v in paths.values()):
            missing = [l for l, v in paths.items() if v is None]
            hh = doc.add_heading(f"{s} — no literal text for: {', '.join(missing)}", level=2)
            hh.runs[0].font.color.rgb = RED
            continue
        wl = {l: words_of(paths[l]) for l in labels}

        # scaffold = median-length config (deterministic, presentation-neutral)
        med = statistics.median_low([len(wl[l]) for l in labels])
        scaffold = sorted([l for l in labels if len(wl[l]) == med])[0]
        base = wl[scaffold]

        maps, ins, changed_all = {}, {}, set()
        for l in labels:
            if l == scaffold:
                continue
            m, i_, ch = map_onto_base(base, wl[l])
            maps[l], ins[l] = m, i_
            changed_all |= ch
        regions = merge_regions(changed_all, len(base))

        rows_data = []
        for (st, en) in regions:
            variants = {l: variant_at(l, scaffold, base, maps, ins, st, en) for l in labels}
            groups = {}
            for l in labels:
                groups.setdefault(variants[l], []).append(l)
            if len(groups) < 2:
                continue          # alignment artefact: everyone actually agrees
            ordered = sorted(groups.items(),
                             key=lambda kv: (-richness(kv[0]), -len(kv[1]), kv[0]))
            # Only claim a "richest" reading when one STRICTLY out-scores the rest.
            # On a tie the ordering is arbitrary, and marking a winner would imply a
            # preference the text doesn't support (e.g. 'twhat' vs 'what').
            strict = len(ordered) > 1 and richness(ordered[0][0]) > richness(ordered[1][0])
            kind = site_kind(list(groups.keys()))
            ctx_pre = " ".join(base[max(0, st - 6):st])
            ctx_post = " ".join(base[en:en + 6])
            rows_data.append((ordered, ctx_pre, ctx_post, strict, kind))

        grand += len(rows_data)
        n_dia = sum(1 for r_ in rows_data if r_[4] == "diacritic")
        n_word = len(rows_data) - n_dia
        dia_total += n_dia
        word_total += n_word
        if not rows_data:
            agree_all += 1
            hh = doc.add_heading(f"{s} — no disagreements", level=2)
            hh.runs[0].font.color.rgb = SLATE
            continue
        hh = doc.add_heading(
            f"{s} — {len(rows_data)} disagreement(s): {n_dia} diacritic, {n_word} wording",
            level=2)
        hh.runs[0].font.color.rgb = SLATE

        t = doc.add_table(rows=1, cols=4)
        t.style = "Light List Accent 1"
        for i, htxt in enumerate(["#", "Reading", "Other reading(s)", "Context"]):
            c = t.rows[0].cells[i]; c.text = ""
            r = c.paragraphs[0].add_run(htxt); r.bold = True; r.font.size = Pt(8)

        for n, (ordered, ctx_pre, ctx_post, strict, kind) in enumerate(rows_data, 1):
            cells = t.add_row().cells
            cells[0].text = ""
            pn = cells[0].paragraphs[0]
            r = pn.add_run(str(n)); r.font.size = Pt(8); r.bold = True
            r = pn.add_run("\n" + kind); r.font.size = Pt(6); r.font.color.rgb = GREY

            # first reading — green ONLY when strictly richest
            text, cfgs = ordered[0]
            cells[1].text = ""
            pr = cells[1].paragraphs[0]
            r = pr.add_run(text if text else "(omitted)")
            r.bold = True; r.font.size = Pt(8)
            if not text:
                r.italic = True; r.font.color.rgb = RED
            elif strict:
                r.font.color.rgb = GREEN
            if strict and text:
                r = pr.add_run("  ◆ richest"); r.font.size = Pt(6); r.font.color.rgb = GREEN
            r = pr.add_run("\n" + "  ".join(cfgs)); r.font.size = Pt(7); r.font.color.rgb = BLUE

            # alternatives
            cells[2].text = ""
            pa = cells[2].paragraphs[0]
            for j, (atext, acfgs) in enumerate(ordered[1:]):
                if j:
                    pa.add_run("\n")
                r = pa.add_run(atext if atext else "(omitted)")
                r.font.size = Pt(8)
                if not atext:
                    r.italic = True; r.font.color.rgb = RED
                r = pa.add_run("  [" + "  ".join(acfgs) + "]")
                r.font.size = Pt(7); r.font.color.rgb = BLUE

            cells[3].text = ""
            pc = cells[3].paragraphs[0]
            r = pc.add_run(("… " if ctx_pre else "") + ctx_pre + " ⟨…⟩ " + ctx_post +
                           (" …" if ctx_post else ""))
            r.font.size = Pt(7); r.font.color.rgb = GREY

    p = doc.add_paragraph()
    r = p.add_run(f"Total disagreement sites across all {len(SCANS)} scans: {grand}")
    r.bold = True
    p.add_run(f"   —   {dia_total} diacritic, {word_total} wording")
    if agree_all:
        p.add_run(f"   ({agree_all} scan(s) with full agreement)")

    out = BASE / "outputs" / args.out
    doc.save(out)
    print(f"saved → {out.relative_to(BASE)}  ({grand} diff sites)")


if __name__ == "__main__":
    main()
