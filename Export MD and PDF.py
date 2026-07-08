#!/usr/bin/env python3
"""
Export MD and PDF.py — optional export step for the Frank Stewart fieldnotes
pipeline, run after "Phase 2 - Normalize Code.py".

Takes a Step-2 output .txt file (literal or normalized, containing
**text**-style bold headings produced from {HEADING}...{/HEADING}) and
exports it to any combination of:

  1. X.md    Markdown — the same text, with single line breaks preserved
             (via trailing hard-break spaces) so it displays correctly in
             any Markdown viewer while keeping **bold** headings intact.

  2. X.pdf   A formatted PDF — the Markdown is rendered to HTML and then to
             PDF, so **bold** headings actually appear bold (not literal
             asterisks) alongside the Unicode diacritics.

  3. X.docx  A Word document — built directly with python-docx, one
             paragraph per blank-line-separated block, one line break per
             '\\n', and a bold run for each **...** span.

Pick which formats to produce with --md / --pdf / --docx (any combination);
if none are given, all three are produced.

Dependencies (not in the standard library):
    pip install markdown weasyprint python-docx

Usage:
    python3 "Export MD and PDF.py" path/to/file.literal.txt
    python3 "Export MD and PDF.py" *.txt --outdir exported/
    python3 "Export MD and PDF.py" file.txt --docx --pdf
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
# markdown + weasyprint are imported lazily inside markdown_to_pdf so that the
# .docx / --combine path works even when those (heavier) libs aren't installed.

FORMATS = ("md", "pdf", "docx")

PAGE_CSS = """
@page {
    size: Letter;
    margin: 1in;
}
body {
    font-family: "Times New Roman", Georgia, serif;
    font-size: 12pt;
    line-height: 1.5;
}
strong {
    font-weight: bold;
}
h1 {
    font-size: 16pt;
    margin-top: 1.2em;
}
.indent {
    text-indent: 2em;
}
"""

HTML_TEMPLATE = """<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>{css}</style>
</head>
<body>
{body}
</body>
</html>
"""

# One block = a **...** bold span or the plain text around it.
_BOLD_SPAN = re.compile(r"(\*\*.*?\*\*)")


def _split_paragraphs(text: str) -> list[dict]:
    """
    Split Step-2 text into paragraph blocks.

    A new paragraph starts wherever a line is indented with leading
    spaces/tabs (the source's typewriter-style paragraph marker) -- this
    is the real paragraph boundary in this text, not blank lines: many
    paragraphs run on from the previous line with only a single '\n'
    between them, indent and all. A blank line also starts a fresh
    paragraph (some section breaks use one instead of/along with an
    indent). Returns a list of {"indent": bool, "lines": [str, ...]}.
    """
    blocks: list[dict | None] = []
    for raw_line in text.split("\n"):
        if raw_line.strip() == "":
            blocks.append(None)
            continue
        stripped = raw_line.lstrip(" \t")
        indented = stripped != raw_line
        starts_new = indented or not blocks or blocks[-1] is None
        if starts_new:
            blocks.append({"indent": indented, "lines": [stripped if indented else raw_line]})
        else:
            blocks[-1]["lines"].append(raw_line)
    return [b for b in blocks if b is not None]


def txt_to_markdown(text: str) -> str:
    """
    Convert Step-2 plain text to Markdown.

    The text already uses **bold** for headings; the only thing missing is
    that a single '\n' is not a paragraph break in Markdown, so bare
    newlines would collapse together. Appending two trailing spaces before
    each newline makes it a Markdown hard line break, preserving the
    original layout.

    Markdown always strips leading whitespace off the first line of a
    paragraph (and treats a 4+-space indent as a code block before it even
    gets that far), so the source's literal paragraph-indent spaces can
    never survive as visible indentation. Instead, each indented paragraph
    is tagged with an attr_list class (`{: .indent}`) and re-indented
    uniformly via CSS in markdown_to_pdf.
    """
    out = []
    for block in _split_paragraphs(text):
        lines = block["lines"]
        rendered = "\n".join(line + "  " for line in lines[:-1]) + lines[-1]
        if block["indent"]:
            rendered += "\n{: .indent}"
        out.append(rendered)
    return "\n\n".join(out)


def markdown_to_pdf(markdown_text: str, pdf_path: Path, title: str) -> None:
    import markdown as md_lib
    from weasyprint import HTML
    body_html = md_lib.markdown(markdown_text, extensions=["attr_list"])
    full_html = HTML_TEMPLATE.format(title=title, css=PAGE_CSS, body=body_html)
    HTML(string=full_html).write_pdf(str(pdf_path))


def _render_text_into_doc(doc, text: str, indent_mult: int = 1) -> None:
    """Append Step-2 text to an open docx Document: a paragraph per blank-line
    block, a line break per '\\n' within a block, and a bold run for each
    **...** span (so {HEADING}-derived headings render bold). {LB} and {INDENT}
    are already rendered to newlines/spaces by Phase 2, so they carry through.

    ``indent_mult`` multiplies each line's leading indentation (the {INDENT}
    spaces), e.g. 2 = twice as deep."""
    blocks = re.split(r"\n\s*\n", text.strip("\n"))
    for block in blocks:
        paragraph = doc.add_paragraph()
        lines = block.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                paragraph.add_run().add_break()
            if indent_mult != 1:
                body = line.lstrip(" ")
                lead = len(line) - len(body)
                line = " " * (lead * indent_mult) + body
            for chunk in _BOLD_SPAN.split(line):
                if not chunk:
                    continue
                if chunk.startswith("**") and chunk.endswith("**") and len(chunk) >= 4:
                    paragraph.add_run(chunk[2:-2]).bold = True
                else:
                    paragraph.add_run(chunk)


def txt_to_docx(text: str, docx_path: Path, indent_mult: int = 1) -> None:
    """Build a single-file .docx from Step-2 text."""
    doc = Document()
    _render_text_into_doc(doc, text, indent_mult)
    doc.save(str(docx_path))


def combine_to_docx(in_paths, docx_path: Path, indent_mult: int = 1) -> None:
    """Combine many Step-2 .txt files into ONE .docx, one transcription per
    page (a page break separates each). Inputs are used in the order given."""
    doc = Document()
    for i, p in enumerate(in_paths):
        if i > 0:
            doc.add_page_break()
        _render_text_into_doc(doc, Path(p).read_text(encoding="utf-8"), indent_mult)
    doc.save(str(docx_path))


def process_file(in_path: Path, outdir: Path | None = None,
                  formats: tuple[str, ...] = FORMATS) -> dict:
    """Export one Step-2 .txt file to any of <stem>.md / .pdf / .docx."""
    in_path = Path(in_path)
    outdir = Path(outdir) if outdir else in_path.parent
    outdir.mkdir(parents=True, exist_ok=True)

    text = in_path.read_text(encoding="utf-8")
    stem = in_path.stem
    written = []

    if "md" in formats or "pdf" in formats:
        markdown_text = txt_to_markdown(text)
        if "md" in formats:
            md_path = outdir / f"{stem}.md"
            md_path.write_text(markdown_text, encoding="utf-8")
            written.append(str(md_path))
        if "pdf" in formats:
            pdf_path = outdir / f"{stem}.pdf"
            markdown_to_pdf(markdown_text, pdf_path, title=stem)
            written.append(str(pdf_path))

    if "docx" in formats:
        docx_path = outdir / f"{stem}.docx"
        txt_to_docx(text, docx_path)
        written.append(str(docx_path))

    return {"source": str(in_path), "outputs": written}


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(
        description="Export Step-2 .txt (literal/normalized) to Markdown, PDF, and/or Word.")
    ap.add_argument("inputs", nargs="+", type=Path,
                     help="Step-2 .txt files to export")
    ap.add_argument("--outdir", type=Path, default=None,
                     help="output directory (default: alongside each input)")
    ap.add_argument("--md", action="store_true", help="export Markdown (.md)")
    ap.add_argument("--pdf", action="store_true", help="export PDF (.pdf)")
    ap.add_argument("--docx", action="store_true", help="export Word (.docx)")
    ap.add_argument("--combine", type=Path, default=None,
                     help="Combine ALL inputs into one .docx at this path, "
                          "one transcription per page (page break between each)")
    ap.add_argument("--indent-mult", type=int, default=1,
                     help="Multiply paragraph indentation depth (e.g. 2 = twice as deep)")
    args = ap.parse_args(argv)

    if args.combine:
        combine_to_docx(args.inputs, args.combine, indent_mult=args.indent_mult)
        print(f"Combined {len(args.inputs)} file(s) -> {args.combine}  (indent x{args.indent_mult})")
        return 0

    requested = tuple(f for f in FORMATS if getattr(args, f))
    formats = requested or FORMATS  # no flags given -> export everything

    for path in args.inputs:
        result = process_file(path, outdir=args.outdir, formats=formats)
        print(f"{path}: wrote {', '.join(Path(o).name for o in result['outputs'])}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
